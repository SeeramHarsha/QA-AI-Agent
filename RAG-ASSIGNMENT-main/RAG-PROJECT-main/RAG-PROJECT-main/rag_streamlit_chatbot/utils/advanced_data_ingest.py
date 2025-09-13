import os
import pdfplumber
import re
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from pdfminer.high_level import extract_text
from pdfminer.layout import LAParams
import arxiv
import logging
from docx import Document as DocxDocument

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_with_pdfplumber(pdf_path):
    """Extract text, tables, and other elements using pdfplumber."""
    content = ""
    tables = []
    images = []
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                # Extract text
                page_text = page.extract_text() or ""
                content += f"\n--- Page {i+1} ---\n{page_text}\n"
                
                # Extract tables
                page_tables = page.extract_tables()
                if page_tables:
                    for j, table in enumerate(page_tables):
                        tables.append({
                            "page": i+1,
                            "table_index": j,
                            "content": table
                        })
                
                # Note image locations (pdfplumber can detect but not extract)
                page_images = page.images
                if page_images:
                    for img in page_images:
                        images.append({
                            "page": i+1,
                            "bbox": img.get("bbox", None)
                        })
                        
        return content, tables, images
    except Exception as e:
        logger.error(f"Error extracting with pdfplumber: {str(e)}")
        raise

def extract_mathematical_equations(text):
    """Extract mathematical equations from text using regex patterns."""
    # Patterns for common mathematical expressions
    patterns = [
        r'\$\$.*?\$\$',  # Display math mode
        r'\$.*?\$',      # Inline math mode
        r'\\begin\{equation\}.*?\\end\{equation\}',
        r'\\begin\{align\}.*?\\end\{align\}',
        r'\\begin\{eqnarray\}.*?\\end\{eqnarray\}',
        r'\\begin\{alignat\}.*?\\end\{alignat\}',
        r'\\begin\{gather\}.*?\\end\{gather\}',
        r'\\begin\{multline\}.*?\\end\{multline\}',
        r'\\[.*?\\]',    # Alternative display math
        r'\(.*?\)',      # Alternative inline math
    ]
    
    equations = []
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.DOTALL)
        for match in matches:
            equation = match.group().strip()
            if len(equation) > 5:  # Filter out very short matches
                equations.append(equation)
    
    return equations

def extract_bibliography(text):
    """Extract bibliography/references section from text."""
    # Common patterns for bibliography sections
    bib_patterns = [
        r'(?:References|Bibliography|Works Cited)[\s\S]*?(?=\n\n|\Z)',
        r'(?:References|Bibliography|Works Cited)\s*\n[\s\S]*?(?=\n\s*\n|\Z)',
        r'(?:References|Bibliography|Works Cited)\s*\n.*?(?=\n[A-Z][a-z]+|\Z)',
        r'\[\d+\].*?(?=\n\[\d+\]|\Z)'  # Reference numbers like [1], [2]
    ]
    
    for pattern in bib_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group().strip()
    
    return None

def extract_figures_and_captions(text):
    """Extract figure references and captions."""
    # Pattern for figure references and captions
    figure_patterns = [
        r'(?:Figure|Fig\.?)\s*\d+[:.]\s*.*?(?=\n\n|\Z)',
        r'(?:Figure|Fig\.?)\s*\d+[a-z]?[:.]\s*.*?(?=\n\n|\Z)',
        r'(?:Figure|Fig\.?)\s*\d+\s*[:.]\s*.*?(?=\n\s*\n|\Z)',
        r'(?:Figure|Fig\.?)\s*\d+\s+\w+.*?(?=\n\n|\Z)'
    ]
    
    figures = []
    for pattern in figure_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
        for match in matches:
            figure_text = match.group().strip()
            if len(figure_text) > 10:  # Filter out very short matches
                figures.append(figure_text)
    
    return figures

def extract_sections(text):
    """Extract section headings and structure."""
    # Pattern for section headings (supports multiple levels)
    section_patterns = [
        r'(\d+(?:\.\d+)*)\s+(.*?)(?=\n\d+(?:\.\d+)*|\Z)',  # Numbered sections
        r'^(Abstract|Introduction|Methods|Methodology|Results|Discussion|Conclusion|References|Appendix)[\s\n]',  # Common section names
        r'([A-Z][a-zA-Z\s]{3,50}?)(?=\n[\s]*[A-Z][a-zA-Z\s]{3,50}?|\Z)'  # Capitalized section titles
    ]
    
    sections = []
    for pattern in section_patterns:
        matches = re.finditer(pattern, text, re.MULTILINE)
        for match in matches:
            if pattern == section_patterns[0]:  # Numbered sections
                sections.append({
                    "number": match.group(1),
                    "title": match.group(2).strip()
                })
            else:
                title = match.group(1).strip()
                if len(title) > 5 and len(title) < 100:  # Filter reasonable titles
                    sections.append({
                        "number": "",
                        "title": title
                    })
    
    return sections

def extract_text_from_docx(docx_path):
    """Extract text content from a DOCX file."""
    try:
        doc = DocxDocument(docx_path)
        content = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += cell.text + "\t"
                content += "\n"
            content += "\n"
        
        return content
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise Exception(f"Error extracting text from DOCX: {str(e)}")


def extract_key_findings(text):
    """Extract key findings and results from research papers."""
    finding_patterns = [
        r'(?:Our results show|We found|The results indicate|We demonstrate|Our findings|Key findings|Main results)[\s\S]*?(?=\n\n|\Z)',
        r'(?:significantly|significant|improvement|better than|outperform|achieve|obtain|yield|produce)[\s\S]{0,200}',
        r'(?:\d+\.?\d*\s*(?:%|percent|times|x|fold))',  # Quantitative results
        r'(?:accuracy|precision|recall|f1|f-score|auc|roc|mae|mse|rms)[\s\S]{0,100}'
    ]
    
    findings = []
    for pattern in finding_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            finding = match.group().strip()
            if len(finding) > 20:  # Filter meaningful findings
                findings.append(finding)
    
    return findings

def extract_abstract(text):
    """Extract abstract from research paper."""
    abstract_patterns = [
        r'(?:Abstract|ABSTRACT)\s*\n([\s\S]*?)(?=\n\n|\Z)',
        r'(?:Abstract|ABSTRACT)\s*\n([\s\S]*?)(?=Introduction|INTRODUCTION|\n1\s)',
        r'(?:Abstract|ABSTRACT)\s*[:.]?\s*([\s\S]{100,1500})'  # Abstract with reasonable length
    ]
    
    for pattern in abstract_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            abstract = match.group(1).strip() if match.lastindex else match.group(0).strip()
            if len(abstract) > 50:  # Filter reasonable abstracts
                return abstract
    
    return None

def load_and_process_advanced(file_path):
    """
    Load and process documents (PDF or DOCX) with enhanced extraction capabilities.
    
    Args:
        file_path (str): Path to the document file (PDF or DOCX)
        
    Returns:
        list: List of Document objects with enhanced metadata
    """
    try:
        # Verify file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Extract content based on file type
        if file_extension == ".pdf":
            # Extract content using multiple methods for PDF
            text_content, tables, images = extract_with_pdfplumber(file_path)
            
            # Extract specialized elements
            equations = extract_mathematical_equations(text_content)
            bibliography = extract_bibliography(text_content)
            figures = extract_figures_and_captions(text_content)
            sections = extract_sections(text_content)
            abstract = extract_abstract(text_content)
            key_findings = extract_key_findings(text_content)
            
            # Create metadata
            metadata = {
                "source": file_path,
                "file_name": os.path.basename(file_path),
                "file_type": "PDF",
                "equations_count": len(equations),
                "tables_count": len(tables),
                "images_count": len(images),
                "figures_count": len(figures),
                "sections_count": len(sections),
                "has_abstract": abstract is not None,
                "has_bibliography": bibliography is not None,
                "key_findings_count": len(key_findings)
            }
            
            # Create document with enhanced content
            enhanced_content = text_content
            
            # Add extracted elements to content if they exist
            if abstract:
                enhanced_content = f"--- ABSTRACT ---\n{abstract}\n\n" + enhanced_content
                
            if key_findings:
                enhanced_content += f"\n\n--- KEY FINDINGS ---\n"
                for i, finding in enumerate(key_findings[:5], 1):  # Limit to top 5 findings
                    enhanced_content += f"{i}. {finding}\n"
                    
            if equations:
                enhanced_content += f"\n\n--- EQUATIONS ---\n"
                for eq in equations[:10]:  # Limit to top 10 equations
                    enhanced_content += f"{eq}\n"
                    
            if bibliography:
                enhanced_content += f"\n\n--- BIBLIOGRAPHY ---\n{bibliography[:2000]}\n"  # Limit length
                
            if figures:
                enhanced_content += f"\n\n--- FIGURES ---\n"
                for fig in figures[:10]:  # Limit to top 10 figures
                    enhanced_content += f"{fig}\n"
                    
        elif file_extension == ".docx":
            # Extract text content from DOCX
            text_content = extract_text_from_docx(file_path)
            
            # For DOCX, we'll use simpler extraction methods
            equations = extract_mathematical_equations(text_content)
            bibliography = extract_bibliography(text_content)
            figures = extract_figures_and_captions(text_content)
            sections = extract_sections(text_content)
            abstract = extract_abstract(text_content)
            key_findings = extract_key_findings(text_content)
            
            # Create metadata
            metadata = {
                "source": file_path,
                "file_name": os.path.basename(file_path),
                "file_type": "DOCX",
                "equations_count": len(equations),
                "figures_count": len(figures),
                "sections_count": len(sections),
                "has_abstract": abstract is not None,
                "has_bibliography": bibliography is not None,
                "key_findings_count": len(key_findings)
            }
            
            # Create document with content
            enhanced_content = text_content
            
            # Add extracted elements to content if they exist
            if abstract:
                enhanced_content = f"--- ABSTRACT ---\n{abstract}\n\n" + enhanced_content
                
            if key_findings:
                enhanced_content += f"\n\n--- KEY FINDINGS ---\n"
                for i, finding in enumerate(key_findings[:5], 1):
                    enhanced_content += f"{i}. {finding}\n"
                    
            if equations:
                enhanced_content += f"\n\n--- EQUATIONS ---\n"
                for eq in equations[:10]:
                    enhanced_content += f"{eq}\n"
                    
            if bibliography:
                enhanced_content += f"\n\n--- BIBLIOGRAPHY ---\n{bibliography[:2000]}\n"
                
            if figures:
                enhanced_content += f"\n\n--- FIGURES ---\n"
                for fig in figures[:10]:
                    enhanced_content += f"{fig}\n"
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Create Document object
        document = Document(
            page_content=enhanced_content,
            metadata=metadata
        )
        
        # Split documents with better chunking strategy
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,  # Increased chunk size for better context
            chunk_overlap=200,  # Increased overlap
            length_function=len,
            separators=["\n\n--- ", "\n\n", "\n", " ", ""]  # Added section separators
        )
        
        split_docs = splitter.split_documents([document])
        
        if not split_docs:
            raise ValueError("Failed to split documents into chunks")
        
        logger.info(f"Processed {file_extension.upper()} with {len(split_docs)} chunks")
        return split_docs
        
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        raise Exception(f"Error processing document: {str(e)}")

def search_arxiv_papers(query, max_results=5):
    """
    Search for academic papers on Arxiv based on a query.
    
    Args:
        query (str): Search query
        max_results (int): Maximum number of results to return
        
    Returns:
        list: List of paper metadata
    """
    try:
        search = arxiv.Search(
            query=query,
            max_results=max_results,
            sort_by=arxiv.SortCriterion.Relevance
        )
        
        papers = []
        for result in search.results():
            papers.append({
                "title": result.title,
                "authors": [author.name for author in result.authors],
                "abstract": result.summary,
                "published": result.published.strftime("%Y-%m-%d"),
                "pdf_url": result.pdf_url,
                "entry_id": result.entry_id
            })
        
        return papers
    except Exception as e:
        logger.error(f"Error searching Arxiv: {str(e)}")
        raise Exception(f"Error searching Arxiv: {str(e)}")

def download_arxiv_paper(paper_entry_id, output_dir="downloads"):
    """
    Download a paper from Arxiv.
    
    Args:
        paper_entry_id (str): Arxiv paper entry ID
        output_dir (str): Directory to save the downloaded paper
        
    Returns:
        str: Path to the downloaded PDF file
    """
    try:
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        # Get paper
        paper = next(arxiv.Search(id_list=[paper_entry_id]).results())
        
        # Download PDF
        filename = f"{paper_entry_id.replace('/', '_')}.pdf"
        filepath = os.path.join(output_dir, filename)
        
        paper.download_pdf(dirpath=output_dir, filename=filename)
        
        return filepath
    except Exception as e:
        logger.error(f"Error downloading paper from Arxiv: {str(e)}")
        raise Exception(f"Error downloading paper from Arxiv: {str(e)}")